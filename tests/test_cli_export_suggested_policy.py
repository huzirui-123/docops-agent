from __future__ import annotations

import importlib
import json
from pathlib import Path
from typing import Any, cast

from typer.testing import CliRunner

from apps.cli.main import app

runner = CliRunner()
yaml = cast(Any, importlib.import_module("yaml"))


def _write_table_template(path: Path) -> None:
    from docx import Document

    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].text = "table"
    document.save(str(path))


def _write_task(path: Path) -> None:
    payload_json = json.dumps({"task_type": "meeting_notice", "payload": {}})
    path.write_text(payload_json, encoding="utf-8")


def test_cli_exports_suggested_policy_even_when_strict_format_fails() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out"
        suggested = root / "suggested.yaml"
        _write_table_template(template)
        _write_task(task)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-mode",
                "strict",
                "--format-baseline",
                "policy",
                "--export-suggested-policy",
                str(suggested),
            ],
        )

        assert result.exit_code == 4
        assert suggested.exists()
        payload = yaml.safe_load(suggested.read_text(encoding="utf-8"))
        assert payload["forbid_tables"] is False
        assert "twips_tolerance" in payload


def test_cli_export_suggested_policy_respects_no_overwrite() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out"
        suggested = root / "suggested.yaml"
        _write_table_template(template)
        _write_task(task)
        sentinel = "marker: keep-me\n"
        suggested.write_text(sentinel, encoding="utf-8")

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--export-suggested-policy",
                str(suggested),
                "--no-overwrite",
            ],
        )

        assert result.exit_code == 1
        assert "--no-overwrite" in result.stdout
        assert suggested.read_text(encoding="utf-8") == sentinel
        assert not (out_dir / "out.docx").exists()
        assert not (out_dir / "out.replace_log.json").exists()
        assert not (out_dir / "out.missing_fields.json").exists()
        assert not (out_dir / "out.format_report.json").exists()

