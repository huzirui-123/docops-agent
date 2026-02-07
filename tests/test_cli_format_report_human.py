from __future__ import annotations

import json
from pathlib import Path

from typer.testing import CliRunner

from apps.cli.main import app

runner = CliRunner()


def _write_table_template(path: Path) -> None:
    from docx import Document

    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].text = "table-body"
    document.save(str(path))


def _write_task(path: Path) -> None:
    payload_json = json.dumps({"task_type": "meeting_notice", "payload": {}})
    path.write_text(payload_json, encoding="utf-8")


def test_cli_format_report_human_prints_summary_by_default() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out_human"
        _write_table_template(template)
        _write_task(task)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-baseline",
                "policy",
            ],
        )

        assert result.exit_code == 0
        assert "format_mode=report" in result.stdout
        assert "result=FAILED" in result.stdout
        assert "issues:" in result.stdout
        assert "suggestion:" in result.stdout


def test_cli_format_report_json_suppresses_human_summary() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out_json"
        _write_table_template(template)
        _write_task(task)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-baseline",
                "policy",
                "--format-report",
                "json",
            ],
        )

        assert result.exit_code == 0
        assert "issues:" not in result.stdout
        assert "suggestion:" not in result.stdout


def test_cli_format_report_strict_failure_exit_code_unchanged() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out_strict"
        _write_table_template(template)
        _write_task(task)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-mode",
                "strict",
                "--format-baseline",
                "policy",
            ],
        )

        assert result.exit_code == 4

