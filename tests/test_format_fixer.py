from __future__ import annotations

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

from core.format.fixer import fix_document
from core.format.models import FormatPolicy
from core.format.validator import validate_document
from core.utils.docx_xml import ensure_paragraph_direct_numpr, get_run_east_asia_font


def _policy(**overrides) -> FormatPolicy:
    payload = {
        "forbid_tables": False,
        "forbid_numpr": True,
        "numpr_direct_only": True,
        "run_font_latin": "BUSINESS_DEFAULT_LATIN",
        "run_font_east_asia": "BUSINESS_DEFAULT_EAST_ASIA",
        "run_size_pt": 12,
        "line_spacing_twips": 360,
        "first_line_indent_twips": 420,
        "twips_tolerance": 20,
        "trim_leading_spaces": True,
        "trim_chars": [" ", "\t", "\u3000"],
    }
    payload.update(overrides)
    return FormatPolicy.model_validate(payload)


def test_fixer_repair_then_validate_passes() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.add_run(" \t")
    paragraph.add_run("\u3000Hello")
    ensure_paragraph_direct_numpr(paragraph)

    before = validate_document(document, _policy(), touched_runs={"p0:r1"})
    assert before.passed is False

    fix_report = fix_document(document, _policy(), touched_runs={"p0:r1"})
    after = validate_document(document, _policy(), touched_runs={"p0:r1"})

    assert fix_report.passed is True
    assert fix_report.fixed_count >= 3
    assert after.passed is True


def test_fixer_does_not_modify_non_touched_runs() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    touched = paragraph.add_run("Touched")
    untouched = paragraph.add_run("Untouched")

    fix_document(document, _policy(), touched_runs={"p0:r0"})

    assert touched.font.name == "BUSINESS_DEFAULT_LATIN"
    assert get_run_east_asia_font(touched) == "BUSINESS_DEFAULT_EAST_ASIA"
    assert untouched.font.name is None
    assert untouched.font.size is None


def test_fixer_table_forbidden_remains_failed() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.add_run("Text")

    report = fix_document(document, _policy(forbid_tables=True), touched_runs={"t0.r0.c0.p0:r0"})

    assert report.passed is False
    assert report.error_count == 1
    assert report.issues[0].code == "TABLE_FORBIDDEN"


def test_fixer_trim_leading_spaces_across_runs() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run(" ")
    paragraph.add_run("\t")
    paragraph.add_run("\u3000Hello")

    fix_document(document, _policy(), touched_runs=set())

    assert paragraph.runs[0].text == ""
    assert paragraph.runs[1].text == ""
    assert paragraph.runs[2].text == "Hello"
