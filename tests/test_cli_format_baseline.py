from __future__ import annotations

import json
from pathlib import Path

from docx.shared import Pt
from typer.testing import CliRunner

from apps.cli.main import app

runner = CliRunner()


def _write_empty_task(path: Path) -> None:
    path.write_text(json.dumps({"task_type": "meeting_notice", "payload": {}}), encoding="utf-8")


def _write_table_template(path: Path) -> None:
    from docx import Document

    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].text = "table-content"
    document.save(str(path))


def _write_indent_template(path: Path, twips: int) -> None:
    from docx import Document

    document = Document()
    paragraph = document.add_paragraph("Indented paragraph")
    paragraph.paragraph_format.first_line_indent = Pt(twips / 20)
    document.save(str(path))


def test_strict_policy_baseline_blocks_table_template() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out"
        _write_table_template(template)
        _write_empty_task(task)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-mode",
                "strict",
                "--format-baseline",
                "policy",
            ],
        )

        assert result.exit_code == 4
        report = json.loads((out_dir / "out.format_report.json").read_text(encoding="utf-8"))
        assert any(issue["code"] == "TABLE_FORBIDDEN" for issue in report["issues"])


def test_strict_template_baseline_allows_table_template() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out"
        _write_table_template(template)
        _write_empty_task(task)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-mode",
                "strict",
                "--format-baseline",
                "template",
            ],
        )

        assert result.exit_code == 0
        report = json.loads((out_dir / "out.format_report.json").read_text(encoding="utf-8"))
        assert not any(issue["code"] == "TABLE_FORBIDDEN" for issue in report["issues"])
        assert report["summary"]["baseline"] == "template"
        assert report["summary"]["effective_policy_overrides"]["forbid_tables"] is False


def test_indent_baseline_switches_from_policy_to_template() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        policy_out = root / "out_policy"
        template_out = root / "out_template"
        _write_indent_template(template, twips=480)
        _write_empty_task(task)

        policy_result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(policy_out),
                "--format-mode",
                "strict",
                "--format-baseline",
                "policy",
            ],
        )

        assert policy_result.exit_code == 4
        policy_report = json.loads(
            (policy_out / "out.format_report.json").read_text(encoding="utf-8")
        )
        assert any(
            issue["code"] == "FIRST_LINE_INDENT_MISMATCH"
            for issue in policy_report["issues"]
        )

        template_result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(template_out),
                "--format-mode",
                "strict",
                "--format-baseline",
                "template",
            ],
        )

        assert template_result.exit_code == 0
        template_report = json.loads(
            (template_out / "out.format_report.json").read_text(encoding="utf-8")
        )
        assert not any(
            issue["code"] == "FIRST_LINE_INDENT_MISMATCH"
            for issue in template_report["issues"]
        )
        assert (
            template_report["summary"]["effective_policy_overrides"]["first_line_indent_twips"]
            == 480
        )
