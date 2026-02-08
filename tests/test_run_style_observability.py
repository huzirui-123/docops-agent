from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt
from typer.testing import CliRunner

from apps.cli.main import app
from core.format.models import FormatPolicy
from core.format.validator import validate_document

runner = CliRunner()


def _policy(**overrides) -> FormatPolicy:
    payload = {
        "forbid_tables": False,
        "forbid_numpr": True,
        "numpr_direct_only": True,
        "run_font_latin": "Calibri",
        "run_font_east_asia": "宋体",
        "run_size_pt": 12,
        "line_spacing_twips": 360,
        "first_line_indent_twips": 420,
        "twips_tolerance": 20,
        "treat_inherited_as_error": False,
        "trim_leading_spaces": True,
        "trim_chars": [" ", "\t", "\u3000"],
    }
    payload.update(overrides)
    return FormatPolicy.model_validate(payload)


def test_template_baseline_none_expected_and_none_rendered_is_consistent() -> None:
    document = Document()
    paragraph = document.add_paragraph("value")
    _ = paragraph.runs[0]

    report = validate_document(
        document,
        _policy(),
        touched_runs={"p0:r0"},
        baseline="template",
        template_run_styles={
            "p0:r0": {
                "latin_font": None,
                "east_asia_font": None,
                "size_pt": None,
            }
        },
    )

    assert report.passed is True
    run_issues = [issue for issue in report.issues if issue.code.startswith("RUN_")]
    assert run_issues == []


def test_policy_baseline_none_rendered_style_is_warn_not_error() -> None:
    document = Document()
    paragraph = document.add_paragraph("value")
    _ = paragraph.runs[0]

    report = validate_document(
        document,
        _policy(),
        touched_runs={"p0:r0"},
        baseline="policy",
    )

    assert report.passed is True
    size_issue = next(issue for issue in report.issues if issue.code == "RUN_SIZE_MISMATCH")
    assert size_issue.severity == "warn"
    assert size_issue.observability == "unknown"


def test_policy_baseline_explicit_mismatch_is_error() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("value")
    run.font.size = Pt(11)

    report = validate_document(
        document,
        _policy(),
        touched_runs={"p0:r0"},
        baseline="policy",
    )

    assert report.passed is False
    size_issue = next(issue for issue in report.issues if issue.code == "RUN_SIZE_MISMATCH")
    assert size_issue.severity == "error"
    assert size_issue.observability == "observed"


def _write_task(path: Path) -> None:
    payload = {"task_type": "meeting_notice", "payload": {"meeting_title": "Kickoff"}}
    path.write_text(json.dumps(payload), encoding="utf-8")


def _write_template(path: Path, *, explicit_size: int | None) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("【MEETING_TITLE】")
    if explicit_size is not None:
        run.font.size = Pt(explicit_size)
    document.save(str(path))


def _write_policy(path: Path, *, treat_inherited_as_error: bool) -> None:
    payload = {
        "forbid_tables": False,
        "forbid_numpr": True,
        "numpr_direct_only": True,
        "run_font_latin": "Calibri",
        "run_font_east_asia": "宋体",
        "run_size_pt": 12,
        "line_spacing_twips": 360,
        "first_line_indent_twips": 420,
        "twips_tolerance": 20,
        "treat_inherited_as_error": treat_inherited_as_error,
        "trim_leading_spaces": True,
        "trim_chars": [" ", "\\t", "\\u3000"],
    }
    lines = [f"{key}: {json.dumps(value, ensure_ascii=False)}" for key, value in payload.items()]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def test_cli_strict_policy_with_inherited_size_warn_does_not_fail() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        policy = root / "policy.yaml"
        out_dir = root / "out"

        _write_template(template, explicit_size=None)
        _write_task(task)
        _write_policy(policy, treat_inherited_as_error=False)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-mode",
                "strict",
                "--format-baseline",
                "policy",
                "--policy",
                str(policy),
            ],
        )

        assert result.exit_code == 0
        payload = json.loads((out_dir / "out.format_report.json").read_text(encoding="utf-8"))
        run_size = [item for item in payload["issues"] if item["code"] == "RUN_SIZE_MISMATCH"]
        assert run_size
        assert all(item["severity"] == "warn" for item in run_size)


def test_cli_strict_policy_with_explicit_size_mismatch_fails() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        policy = root / "policy.yaml"
        out_dir = root / "out"

        _write_template(template, explicit_size=11)
        _write_task(task)
        _write_policy(policy, treat_inherited_as_error=False)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-mode",
                "strict",
                "--format-baseline",
                "policy",
                "--policy",
                str(policy),
            ],
        )

        assert result.exit_code == 4
        payload = json.loads((out_dir / "out.format_report.json").read_text(encoding="utf-8"))
        run_size = [item for item in payload["issues"] if item["code"] == "RUN_SIZE_MISMATCH"]
        assert run_size
        assert any(item["severity"] == "error" for item in run_size)
