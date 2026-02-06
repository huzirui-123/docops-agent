from __future__ import annotations

import pytest
from docx import Document

from core.templates.placeholder_parser import parse_placeholders
from core.utils.errors import TemplateError


def test_parse_single_placeholder_in_single_run() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("AA【NAME】BB")

    result = parse_placeholders(document)

    assert result.fields == ["NAME"]
    assert len(result.occurrences) == 1
    assert result.occurrences[0].run_id == "p0:r0"
    assert result.occurrences[0].start == 2
    assert result.occurrences[0].end == 8
    assert not result.unsupported


def test_parse_multiple_placeholders_in_single_run() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("【A1】X【B_2】")

    result = parse_placeholders(document)

    assert result.fields == ["A1", "B_2"]
    assert len(result.occurrences) == 2
    assert not result.unsupported


def test_parse_cross_run_placeholder_records_unsupported() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("AA【NA")
    paragraph.add_run("ME】BB")

    result = parse_placeholders(document)

    assert result.fields == []
    assert not result.occurrences
    assert len(result.unsupported) == 1
    assert result.unsupported[0].kind == "cross_run"


def test_parse_cross_run_placeholder_raises_in_strict_mode() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("【AB")
    paragraph.add_run("CD】")

    with pytest.raises(TemplateError):
        parse_placeholders(document, strict=True)


def test_parse_cross_three_runs_records_unsupported() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("【A")
    paragraph.add_run("BC")
    paragraph.add_run("D】")

    result = parse_placeholders(document)

    assert len(result.unsupported) == 1
    assert result.unsupported[0].kind == "cross_run"


def test_parse_table_cell_placeholder() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.text = "X【CELL】Y"

    result = parse_placeholders(document)

    assert result.fields == ["CELL"]
    assert len(result.occurrences) == 1
    assert result.occurrences[0].run_id == "t0.r0.c0.p0:r0"
    assert not result.unsupported


def test_parse_table_cell_cross_run_records_unsupported() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.add_run("【CE")
    paragraph.add_run("LL】")

    result = parse_placeholders(document)

    assert len(result.unsupported) == 1
    assert result.unsupported[0].kind == "cross_run"
    assert result.unsupported[0].run_id == "t0.r0.c0.p0:r0"


def test_parse_invalid_field_name_records_unsupported() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("【bad-name】")

    result = parse_placeholders(document)

    assert not result.occurrences
    assert len(result.unsupported) == 1
    assert result.unsupported[0].kind == "invalid_format"


def test_parse_unclosed_open_bracket_records_unsupported() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Hello【FIELD")

    result = parse_placeholders(document)

    assert len(result.unsupported) == 1
    assert result.unsupported[0].kind == "unclosed_bracket"


def test_parse_stray_close_bracket_records_unsupported() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Hello】World")

    result = parse_placeholders(document)

    assert len(result.unsupported) == 1
    assert result.unsupported[0].kind == "stray_close"


def test_parse_mixed_valid_and_invalid_tokens() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("【OK】 and 【bad】")

    result = parse_placeholders(document)

    assert result.fields == ["OK"]
    assert len(result.occurrences) == 1
    assert len(result.unsupported) == 1
    assert result.unsupported[0].kind == "invalid_format"


def test_parse_adjacent_text_and_placeholder() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Hello【USER_1】World")

    result = parse_placeholders(document)

    assert result.fields == ["USER_1"]
    assert len(result.occurrences) == 1
    assert result.occurrences[0].start == 5
    assert result.occurrences[0].end == 13


def test_parse_header_footer_are_ignored() -> None:
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "【HEADER】"
    section.footer.paragraphs[0].text = "【FOOTER】"

    result = parse_placeholders(document, strict=True)

    assert result.fields == []
    assert result.occurrences == []
    assert result.unsupported == []


def test_parse_strict_mode_raises_on_unclosed_and_stray() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("A】B【FIELD")

    with pytest.raises(TemplateError) as exc_info:
        parse_placeholders(document, strict=True)

    assert exc_info.value.result is not None
    kinds = [item.kind for item in exc_info.value.result.unsupported]
    assert "stray_close" in kinds
    assert "unclosed_bracket" in kinds
