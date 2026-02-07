from __future__ import annotations

import json
from pathlib import Path

from typer.testing import CliRunner

from apps.cli.main import app

runner = CliRunner()


def _write_table_docx(path: Path) -> None:
    from docx import Document

    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].text = "table-body"
    document.save(str(path))


def _write_task(path: Path) -> None:
    payload_json = json.dumps({"task_type": "meeting_notice", "payload": {}})
    path.write_text(payload_json, encoding="utf-8")


def test_cli_default_report_mode_does_not_block_format_failures() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out"
        _write_table_docx(template)
        _write_task(task)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
            ],
        )

        assert result.exit_code == 0
        assert "WARNING(format): format issues detected" in result.stdout
        report = json.loads((out_dir / "out.format_report.json").read_text(encoding="utf-8"))
        assert report["passed"] is False
        assert any(issue["code"] == "TABLE_FORBIDDEN" for issue in report["issues"])


def test_cli_strict_mode_blocks_with_exit_4() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out"
        _write_table_docx(template)
        _write_task(task)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-mode",
                "strict",
            ],
        )

        assert result.exit_code == 4
        report = json.loads((out_dir / "out.format_report.json").read_text(encoding="utf-8"))
        assert any(issue["code"] == "TABLE_FORBIDDEN" for issue in report["issues"])


def test_cli_off_mode_skips_format_steps() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out"
        _write_table_docx(template)
        _write_task(task)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-mode",
                "off",
            ],
        )

        assert result.exit_code == 0
        assert "INFO: format validation skipped" in result.stdout
        report = json.loads((out_dir / "out.format_report.json").read_text(encoding="utf-8"))
        assert report["passed"] is True
        assert report["issues"] == []
        assert report["summary"]["skipped"] is True
        assert report["summary"]["mode"] == "off"
