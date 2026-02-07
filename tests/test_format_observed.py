from __future__ import annotations

from docx import Document
from docx.shared import Pt

from core.format.observed import diff_observed, observe_document
from core.utils.docx_xml import ensure_paragraph_direct_numpr


def test_observe_document_collects_expected_histograms() -> None:
    document = Document()
    paragraph = document.add_paragraph("Alpha")
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.runs[0].font.name = "Calibri"
    ensure_paragraph_direct_numpr(paragraph)

    plain = document.add_paragraph("Beta")
    plain.add_run(" tail")

    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].text = "Cell"

    observed = observe_document(document)

    assert observed.has_tables is True
    assert observed.has_numpr is True
    assert observed.first_line_indent_twips_hist["420"] >= 1
    assert observed.first_line_indent_twips_hist["none"] >= 1
    assert observed.run_font_latin_hist["Calibri"] >= 1
    assert observed.run_font_latin_hist["none"] >= 1
    assert observed.run_font_east_asia_hist["unknown"] >= 1


def test_diff_observed_calculates_indent_delta() -> None:
    template_document = Document()
    template_document.add_paragraph("One")

    rendered_document = Document()
    paragraph = rendered_document.add_paragraph("One")
    paragraph.paragraph_format.first_line_indent = Pt(24)

    template_observed = observe_document(template_document)
    rendered_observed = observe_document(rendered_document)
    delta = diff_observed(template_observed, rendered_observed)

    assert delta.has_tables_changed is False
    assert delta.has_numpr_changed is False
    assert delta.first_line_indent_twips_hist_delta["480"] == 1
    assert delta.first_line_indent_twips_hist_delta["none"] == -1
