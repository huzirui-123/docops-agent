from __future__ import annotations

from docx import Document

from core.format.models import FormatPolicy
from core.orchestrator.pipeline import run_task
from core.skills.base import Skill
from core.skills.models import SkillResult, TaskSpec
from core.utils.errors import MissingRequiredFieldsError


class StaticSkill(Skill):
    name = "static"

    def __init__(self, result: SkillResult) -> None:
        self._result = result

    def build_fields(self, task: TaskSpec) -> SkillResult:
        return self._result


def _policy(**overrides) -> FormatPolicy:
    payload = {
        "forbid_tables": False,
        "forbid_numpr": True,
        "numpr_direct_only": True,
        "run_font_latin": "BUSINESS_DEFAULT_LATIN",
        "run_font_east_asia": "BUSINESS_DEFAULT_EAST_ASIA",
        "run_size_pt": 12,
        "line_spacing_twips": 360,
        "first_line_indent_twips": 420,
        "twips_tolerance": 20,
        "trim_leading_spaces": True,
        "trim_chars": [" ", "\t", "\u3000"],
    }
    payload.update(overrides)
    return FormatPolicy.model_validate(payload)


def test_pipeline_warn_path_exposes_unsupported_summary() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("【OK】 and ")
    paragraph.add_run("【BA")
    paragraph.add_run("D】")

    skill = StaticSkill(
        SkillResult(field_values={"OK": "DONE"}, required_fields={"OK"}, optional_fields=set())
    )

    output = run_task(
        task_spec=TaskSpec(task_type="meeting_notice", payload={}),
        template_document=document,
        skill=skill,
        policy=_policy(),
        unsupported_mode="warn",
    )

    assert output.replace_report.summary.had_unsupported is True
    assert output.replace_report.summary.unsupported_count == 1
    assert output.replace_report.summary.unsupported_mode == "warn"


def test_pipeline_missing_required_uses_template_fields_boundary() -> None:
    document = Document()
    document.add_paragraph("【A】")

    skill = StaticSkill(
        SkillResult(
            field_values={},
            required_fields={"A", "NOT_IN_TEMPLATE"},
            optional_fields=set(),
        )
    )

    try:
        run_task(
            task_spec=TaskSpec(task_type="meeting_notice", payload={}),
            template_document=document,
            skill=skill,
            policy=_policy(),
        )
    except MissingRequiredFieldsError as exc:
        assert exc.missing_required == ["A"]
        assert exc.render_output is not None
        assert exc.render_output.template_fields == ["A"]
    else:
        raise AssertionError("Expected MissingRequiredFieldsError")


def test_pipeline_success_path_passes_formatter() -> None:
    document = Document()
    document.add_paragraph("【A】")

    skill = StaticSkill(
        SkillResult(field_values={"A": "ok"}, required_fields={"A"}, optional_fields=set())
    )

    output = run_task(
        task_spec=TaskSpec(task_type="meeting_notice", payload={}),
        template_document=document,
        skill=skill,
        policy=_policy(),
    )

    assert output.document.paragraphs[0].runs[0].text == "ok"
    assert output.format_report.passed is True


def test_pipeline_returns_failed_format_report_for_non_fixable_issue() -> None:
    document = Document()
    document.add_table(rows=1, cols=1)

    skill = StaticSkill(SkillResult(field_values={}, required_fields=set(), optional_fields=set()))

    output = run_task(
        task_spec=TaskSpec(task_type="meeting_notice", payload={}),
        template_document=document,
        skill=skill,
        policy=_policy(forbid_tables=True),
        unsupported_mode="warn",
    )

    assert output.format_report.passed is False
    assert any(issue.code == "TABLE_FORBIDDEN" for issue in output.format_report.issues)
