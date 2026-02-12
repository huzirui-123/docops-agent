from __future__ import annotations

import io
import json

import httpx
import pytest
from docx import Document

from apps.api.main import app


def _build_single_run_docx(placeholder: str) -> bytes:
    document = Document()
    document.add_paragraph(placeholder)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _build_cross_run_docx() -> bytes:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("【MEETING_")
    paragraph.add_run("TITLE】")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.mark.anyio
async def test_precheck_success_returns_expected_exit_code_zero() -> None:
    task_bytes = json.dumps(
        {"task_type": "meeting_notice", "payload": {"meeting_title": "Weekly Sync"}}
    ).encode("utf-8")
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
        response = await client.post(
            "/v1/precheck",
            files={
                "template": (
                    "template.docx",
                    _build_single_run_docx("【MEETING_TITLE】"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                "task": ("task.json", task_bytes, "application/json"),
            },
            data={"skill": "meeting_notice"},
        )

    assert response.status_code == 200
    request_id = response.headers["X-Docops-Request-Id"]
    payload = response.json()
    assert payload["ok"] is True
    assert payload["expected_exit_code"] == 0
    assert payload["request_id"] == request_id
    assert payload["summary"]["unsupported_count"] == 0
    assert payload["missing_required"] == []


@pytest.mark.anyio
async def test_precheck_missing_required_returns_expected_exit_code_two() -> None:
    task_bytes = json.dumps({"task_type": "meeting_notice", "payload": {}}).encode("utf-8")
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
        response = await client.post(
            "/v1/precheck",
            files={
                "template": (
                    "template.docx",
                    _build_single_run_docx("【MEETING_TITLE】"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                "task": ("task.json", task_bytes, "application/json"),
            },
            data={"skill": "meeting_notice"},
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["ok"] is False
    assert payload["expected_exit_code"] == 2
    assert "MEETING_TITLE" in payload["missing_required"]
    assert payload["summary"]["missing_required_count"] >= 1


@pytest.mark.anyio
async def test_precheck_unsupported_placeholder_returns_expected_exit_code_three() -> None:
    task_bytes = json.dumps(
        {"task_type": "meeting_notice", "payload": {"meeting_title": "Weekly Sync"}}
    ).encode("utf-8")
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
        response = await client.post(
            "/v1/precheck",
            files={
                "template": (
                    "template.docx",
                    _build_cross_run_docx(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                "task": ("task.json", task_bytes, "application/json"),
            },
            data={"skill": "meeting_notice"},
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["ok"] is False
    assert payload["expected_exit_code"] == 3
    assert payload["summary"]["unsupported_count"] >= 1


@pytest.mark.anyio
async def test_precheck_rejects_skill_task_type_mismatch() -> None:
    task_bytes = json.dumps({"task_type": "meeting_notice", "payload": {}}).encode("utf-8")
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
        response = await client.post(
            "/v1/precheck",
            files={
                "template": (
                    "template.docx",
                    _build_single_run_docx("【MEETING_TITLE】"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                "task": ("task.json", task_bytes, "application/json"),
            },
            data={"skill": "training_notice"},
        )

    assert response.status_code == 400
    assert response.headers["X-Docops-Request-Id"]
    payload = response.json()
    assert payload["error_code"] == "INVALID_ARGUMENT_CONFLICT"
    assert payload["detail"]["request_id"] == response.headers["X-Docops-Request-Id"]
