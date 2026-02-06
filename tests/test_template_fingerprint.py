from __future__ import annotations

from docx import Document

from core.templates.models import ParseResult, UnsupportedOccurrence
from core.templates.template_fingerprint import compute_template_fingerprint


def test_fingerprint_is_stable_for_same_document() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Hello 【NAME】")

    first = compute_template_fingerprint(document)
    second = compute_template_fingerprint(document)

    assert first == second


def test_fingerprint_ignores_run_split_for_same_full_text() -> None:
    doc_a = Document()
    p_a = doc_a.add_paragraph()
    p_a.add_run("Hello ")
    p_a.add_run("World")

    doc_b = Document()
    p_b = doc_b.add_paragraph()
    p_b.add_run("Hel")
    p_b.add_run("lo W")
    p_b.add_run("orld")

    assert compute_template_fingerprint(doc_a) == compute_template_fingerprint(doc_b)


def test_fingerprint_changes_when_body_text_changes() -> None:
    doc_a = Document()
    doc_a.add_paragraph("Alpha")

    doc_b = Document()
    doc_b.add_paragraph("Beta")

    assert compute_template_fingerprint(doc_a) != compute_template_fingerprint(doc_b)


def test_fingerprint_changes_when_table_text_changes() -> None:
    doc_a = Document()
    table_a = doc_a.add_table(rows=1, cols=1)
    table_a.cell(0, 0).paragraphs[0].text = "Cell A"

    doc_b = Document()
    table_b = doc_b.add_table(rows=1, cols=1)
    table_b.cell(0, 0).paragraphs[0].text = "Cell B"

    assert compute_template_fingerprint(doc_a) != compute_template_fingerprint(doc_b)


def test_fingerprint_ignores_header_footer_changes() -> None:
    doc_a = Document()
    doc_a.add_paragraph("Body")
    doc_a.sections[0].header.paragraphs[0].text = "Header A"
    doc_a.sections[0].footer.paragraphs[0].text = "Footer A"

    doc_b = Document()
    doc_b.add_paragraph("Body")
    doc_b.sections[0].header.paragraphs[0].text = "Header B"
    doc_b.sections[0].footer.paragraphs[0].text = "Footer B"

    assert compute_template_fingerprint(doc_a) == compute_template_fingerprint(doc_b)


def test_fingerprint_stable_with_valid_and_unsupported_mix() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("X【OK】Y")
    paragraph.add_run("【ba")
    paragraph.add_run("d】")

    first = compute_template_fingerprint(document)
    second = compute_template_fingerprint(document)

    assert first == second


def test_fingerprint_unsupported_without_position_has_stable_sort(monkeypatch) -> None:
    document = Document()
    document.add_paragraph("Body")

    call_index = {"value": 0}

    def fake_parse_placeholders(_document, strict=False):
        call_index["value"] += 1
        if call_index["value"] % 2 == 1:
            unsupported = [
                UnsupportedOccurrence(
                    kind="invalid_format",
                    text="【bad2】",
                    run_id=None,
                    start=None,
                    end=None,
                ),
                UnsupportedOccurrence(
                    kind="invalid_format",
                    text="【bad1】",
                    run_id=None,
                    start=None,
                    end=None,
                ),
            ]
        else:
            unsupported = [
                UnsupportedOccurrence(
                    kind="invalid_format",
                    text="【bad1】",
                    run_id=None,
                    start=None,
                    end=None,
                ),
                UnsupportedOccurrence(
                    kind="invalid_format",
                    text="【bad2】",
                    run_id=None,
                    start=None,
                    end=None,
                ),
            ]
        return ParseResult(fields=[], occurrences=[], unsupported=unsupported)

    monkeypatch.setattr(
        "core.templates.template_fingerprint.parse_placeholders", fake_parse_placeholders
    )

    first = compute_template_fingerprint(document)
    second = compute_template_fingerprint(document)

    assert first == second
