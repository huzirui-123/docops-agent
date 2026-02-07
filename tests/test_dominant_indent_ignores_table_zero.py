from __future__ import annotations

from docx import Document
from docx.shared import Pt

from core.format.models import FormatPolicy
from core.orchestrator.pipeline import run_task
from core.skills.meeting_notice import MeetingNoticeSkill
from core.skills.models import TaskSpec


def _policy() -> FormatPolicy:
    return FormatPolicy.model_validate(
        {
            "forbid_tables": False,
            "forbid_numpr": True,
            "numpr_direct_only": True,
            "run_font_latin": "Calibri",
            "run_font_east_asia": "宋体",
            "run_size_pt": 12,
            "line_spacing_twips": 360,
            "first_line_indent_twips": 420,
            "twips_tolerance": 20,
            "trim_leading_spaces": True,
            "trim_chars": [" ", "\t", "\u3000"],
        }
    )


def test_template_dominant_indent_prefers_body_over_table_zero() -> None:
    document = Document()
    body_paragraph = document.add_paragraph("Body line")
    body_paragraph.paragraph_format.first_line_indent = Pt(24)  # 480 twips

    table = document.add_table(rows=3, cols=2)
    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            paragraph.text = "cell"
            paragraph.paragraph_format.first_line_indent = Pt(0)

    output = run_task(
        task_spec=TaskSpec(task_type="meeting_notice", payload={}),
        template_document=document,
        skill=MeetingNoticeSkill(),
        policy=_policy(),
        format_mode="strict",
        format_baseline="template",
    )

    assert output.format_report.summary is not None
    assert output.format_report.summary.effective_policy_overrides["first_line_indent_twips"] == 480
    assert any(
        issue.code == "FIRST_LINE_INDENT_MISMATCH" and issue.expected == 480
        for issue in output.format_report.issues
    )
