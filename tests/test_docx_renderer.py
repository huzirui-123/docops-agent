from __future__ import annotations

import pytest
from docx import Document

from core.render.docx_renderer import render_docx
from core.skills.models import SkillResult
from core.utils.errors import TemplateError


def test_render_output_contains_parse_result_and_template_fields() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Title: 【MEETING_TITLE】")

    output = render_docx(
        document=document,
        skill_result=SkillResult(
            field_values={"MEETING_TITLE": "Kickoff"},
            required_fields={"MEETING_TITLE"},
            optional_fields=set(),
        ),
    )

    assert output.template_fields == ["MEETING_TITLE"]
    assert output.parse_result.fields == ["MEETING_TITLE"]


def test_missing_fields_respect_template_field_boundary() -> None:
    document = Document()
    document.add_paragraph("【A】")

    output = render_docx(
        document=document,
        skill_result=SkillResult(
            field_values={},
            required_fields={"A", "B"},
            optional_fields={"C"},
        ),
        unsupported_mode="warn",
    )

    assert output.missing_fields.missing_required == ["A"]
    assert output.missing_fields.missing_optional == []


def test_replacement_orders_same_run_occurrences_by_desc_start() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("【A】【B】")

    output = render_docx(
        document=document,
        skill_result=SkillResult(
            field_values={"A": "LONGA", "B": "X"},
            required_fields={"A", "B"},
            optional_fields=set(),
        ),
    )

    assert output.document.paragraphs[0].runs[0].text == "LONGAX"
    assert output.replace_report.template_run_styles["p0:r0"].size_pt is None


def test_warn_mode_summary_marks_unsupported_metadata() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("【OK】 and ")
    paragraph.add_run("【BA")
    paragraph.add_run("D】")

    output = render_docx(
        document=document,
        skill_result=SkillResult(
            field_values={"OK": "VALUE"},
            required_fields={"OK"},
            optional_fields=set(),
        ),
        unsupported_mode="warn",
    )

    summary = output.replace_report.summary
    assert summary.had_unsupported is True
    assert summary.unsupported_count == 1
    assert summary.unsupported_mode == "warn"
    assert output.document.paragraphs[0].runs[0].text == "VALUE and "


def test_error_mode_raises_with_replace_report() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("【BA")
    paragraph.add_run("D】")

    with pytest.raises(TemplateError) as exc_info:
        render_docx(
            document=document,
            skill_result=SkillResult(field_values={}, required_fields=set(), optional_fields=set()),
            unsupported_mode="error",
        )

    replace_report = exc_info.value.replace_report
    assert replace_report is not None
    assert replace_report.summary.had_unsupported is True
    assert replace_report.summary.unsupported_count == 1
    assert replace_report.summary.unsupported_mode == "error"
