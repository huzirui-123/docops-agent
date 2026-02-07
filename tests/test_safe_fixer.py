from __future__ import annotations

from docx import Document
from docx.shared import Pt

from core.format.models import FormatPolicy
from core.format.safe_fixer import safe_fix_document
from core.utils.docx_xml import get_first_line_indent_twips


def _policy() -> FormatPolicy:
    return FormatPolicy.model_validate(
        {
            "forbid_tables": True,
            "forbid_numpr": True,
            "numpr_direct_only": True,
            "run_font_latin": "Calibri",
            "run_font_east_asia": "宋体",
            "run_size_pt": 12,
            "line_spacing_twips": 360,
            "first_line_indent_twips": 420,
            "twips_tolerance": 20,
            "trim_leading_spaces": True,
            "trim_chars": [" ", "\t", "\u3000"],
        }
    )


def test_safe_fixer_updates_only_touched_body_paragraph() -> None:
    document = Document()
    touched_paragraph = document.add_paragraph("Touched")
    untouched_paragraph = document.add_paragraph("Untouched")
    touched_paragraph.paragraph_format.first_line_indent = Pt(24)  # 480
    untouched_paragraph.paragraph_format.first_line_indent = Pt(24)  # 480

    changes = safe_fix_document(document, _policy(), touched_runs={"p0:r0"})

    assert get_first_line_indent_twips(touched_paragraph) == 420
    assert get_first_line_indent_twips(untouched_paragraph) == 480
    assert any(
        change["paragraph_path"] == "p0"
        and change["field"] == "first_line_indent_twips"
        and change["before"] == 480
        and change["after"] == 420
        for change in changes
    )


def test_safe_fixer_does_not_modify_table_paragraphs() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    cell_paragraph = table.cell(0, 0).paragraphs[0]
    cell_paragraph.text = "Cell"
    cell_paragraph.paragraph_format.first_line_indent = Pt(24)  # 480

    changes = safe_fix_document(document, _policy(), touched_runs={"t0.r0.c0.p0:r0"})

    assert get_first_line_indent_twips(cell_paragraph) == 480
    assert changes == []

