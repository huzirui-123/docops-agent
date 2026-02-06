from __future__ import annotations

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

from core.format.models import FormatPolicy
from core.format.validator import validate_document
from core.utils.docx_xml import ensure_paragraph_direct_numpr


def _policy(**overrides) -> FormatPolicy:
    payload = {
        "forbid_tables": True,
        "forbid_numpr": True,
        "numpr_direct_only": True,
        "run_font_latin": "BUSINESS_DEFAULT_LATIN",
        "run_font_east_asia": "BUSINESS_DEFAULT_EAST_ASIA",
        "run_size_pt": 12,
        "line_spacing_twips": 360,
        "first_line_indent_twips": 420,
        "twips_tolerance": 20,
        "trim_leading_spaces": True,
        "trim_chars": [" ", "\t", "\u3000"],
    }
    payload.update(overrides)
    return FormatPolicy.model_validate(payload)


def test_validator_table_forbidden_skips_table_cell_checks() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.add_run("Text")
    ensure_paragraph_direct_numpr(paragraph)

    report = validate_document(
        document, _policy(forbid_tables=True), touched_runs={"t0.r0.c0.p0:r0"}
    )

    assert report.passed is False
    assert report.error_count == 1
    assert [issue.code for issue in report.issues] == ["TABLE_FORBIDDEN"]


def test_validator_table_checks_enabled_when_not_forbidden() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.add_run("Text")
    ensure_paragraph_direct_numpr(paragraph)

    report = validate_document(document, _policy(forbid_tables=False), touched_runs=set())

    assert any(issue.code == "NUMPR_PRESENT" for issue in report.issues)


def test_validator_line_indent_tolerance() -> None:
    document = Document()
    paragraph = document.add_paragraph("Body")
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(21)

    pass_report = validate_document(document, _policy(twips_tolerance=0), touched_runs=set())
    fail_report = validate_document(
        document,
        _policy(line_spacing_twips=300, first_line_indent_twips=300, twips_tolerance=10),
        touched_runs=set(),
    )

    assert pass_report.passed is True
    assert any(issue.code == "LINE_SPACING_MISMATCH" for issue in fail_report.issues)
    assert any(issue.code == "FIRST_LINE_INDENT_MISMATCH" for issue in fail_report.issues)


def test_validator_line_spacing_auto_is_ignored() -> None:
    document = Document()
    paragraph = document.add_paragraph("Body")
    paragraph.paragraph_format.line_spacing = 1.5

    report = validate_document(document, _policy(line_spacing_twips=360, twips_tolerance=0), set())

    assert not any(issue.code == "LINE_SPACING_MISMATCH" for issue in report.issues)


def test_validator_checks_only_touched_run_font_and_size() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("A")
    paragraph.add_run("B")

    report = validate_document(document, _policy(), touched_runs={"p0:r0"})

    run_issue_ids = [issue.run_id for issue in report.issues if issue.code.startswith("RUN_")]
    assert run_issue_ids == ["p0:r0", "p0:r0"]
