from __future__ import annotations

from docx import Document
from docx.shared import Pt

from core.format.models import FormatPolicy
from core.format.validator import validate_document
from core.orchestrator.pipeline import run_task
from core.skills.meeting_notice import MeetingNoticeSkill
from core.skills.models import TaskSpec


def _policy(**overrides) -> FormatPolicy:
    payload = {
        "forbid_tables": True,
        "forbid_numpr": True,
        "numpr_direct_only": True,
        "run_font_latin": "Calibri",
        "run_font_east_asia": "宋体",
        "run_size_pt": 12,
        "line_spacing_twips": 360,
        "first_line_indent_twips": 420,
        "twips_tolerance": 20,
        "trim_leading_spaces": True,
        "trim_chars": [" ", "\t", "\u3000"],
    }
    payload.update(overrides)
    return FormatPolicy.model_validate(payload)


def test_validator_populates_structured_indent_issue_fields() -> None:
    document = Document()
    paragraph = document.add_paragraph("Indented")
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 480 twips

    report = validate_document(document, _policy(), touched_runs=set())
    issue = next(item for item in report.issues if item.code == "FIRST_LINE_INDENT_MISMATCH")

    assert issue.expected == 420
    assert issue.actual == 480
    assert issue.tolerance == 20


def test_pipeline_report_and_off_diagnostics_behavior() -> None:
    document = Document()
    document.add_table(rows=1, cols=1)

    report_output = run_task(
        task_spec=TaskSpec(task_type="meeting_notice", payload={}),
        template_document=document,
        skill=MeetingNoticeSkill(),
        policy=_policy(),
        format_mode="report",
        format_baseline="policy",
    )

    assert report_output.format_report.summary is not None
    diagnostics = report_output.format_report.summary.diagnostics
    assert diagnostics is not None
    assert diagnostics["by_code"]["TABLE_FORBIDDEN"]["count"] >= 1
    assert diagnostics["by_code"]["TABLE_FORBIDDEN"]["examples"]
    assert diagnostics["by_code"]["TABLE_FORBIDDEN"]["suggestions"]

    off_document = Document()
    off_document.add_table(rows=1, cols=1)
    off_output = run_task(
        task_spec=TaskSpec(task_type="meeting_notice", payload={}),
        template_document=off_document,
        skill=MeetingNoticeSkill(),
        policy=_policy(),
        format_mode="off",
    )

    assert off_output.format_report.summary is not None
    assert off_output.format_report.summary.skipped is True
    assert off_output.format_report.summary.diagnostics is None
