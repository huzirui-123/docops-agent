from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt
from typer.testing import CliRunner

from apps.cli.main import app

runner = CliRunner()


def _write_template(path: Path, *, explicit_size: int | None) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("【MEETING_TITLE】")
    if explicit_size is not None:
        run.font.size = Pt(explicit_size)
    document.save(str(path))


def _write_task(path: Path) -> None:
    payload = {"task_type": "meeting_notice", "payload": {"meeting_title": "Kickoff"}}
    path.write_text(json.dumps(payload), encoding="utf-8")


def _write_policy(path: Path, *, treat_inherited_as_error: bool = False) -> None:
    payload = {
        "forbid_tables": False,
        "forbid_numpr": True,
        "numpr_direct_only": True,
        "run_font_latin": "Calibri",
        "run_font_east_asia": "宋体",
        "run_size_pt": 12,
        "line_spacing_twips": 360,
        "first_line_indent_twips": 420,
        "twips_tolerance": 20,
        "treat_inherited_as_error": treat_inherited_as_error,
        "trim_leading_spaces": True,
        "trim_chars": [" ", "\\t", "\\u3000"],
    }
    lines = [f"{key}: {json.dumps(value, ensure_ascii=False)}" for key, value in payload.items()]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def test_warn_only_report_shows_summary_warnings_without_warning_line() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        policy = root / "policy.yaml"
        out_dir = root / "out"

        _write_template(template, explicit_size=None)
        _write_task(task)
        _write_policy(policy, treat_inherited_as_error=False)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-mode",
                "report",
                "--format-baseline",
                "policy",
                "--format-report",
                "human",
                "--policy",
                str(policy),
            ],
        )

        assert result.exit_code == 0
        assert "warnings:" in result.stdout
        assert "warnings: none" not in result.stdout
        assert "WARNING(format):" not in result.stdout
        assert "next_cmd:" in result.stdout
        assert "--format-baseline template" in result.stdout


def test_warn_only_strict_still_returns_zero() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        policy = root / "policy.yaml"
        out_dir = root / "out"

        _write_template(template, explicit_size=None)
        _write_task(task)
        _write_policy(policy, treat_inherited_as_error=False)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-mode",
                "strict",
                "--format-baseline",
                "policy",
                "--policy",
                str(policy),
            ],
        )

        assert result.exit_code == 0
        assert "warnings:" in result.stdout
        assert "warnings: none" not in result.stdout
        assert "WARNING(format):" not in result.stdout


def test_error_strict_prints_warning_and_fails() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        policy = root / "policy.yaml"
        out_dir = root / "out"

        _write_template(template, explicit_size=11)
        _write_task(task)
        _write_policy(policy, treat_inherited_as_error=False)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-mode",
                "strict",
                "--format-baseline",
                "policy",
                "--format-report",
                "human",
                "--policy",
                str(policy),
            ],
        )

        assert result.exit_code == 4
        assert "WARNING(format): format issues detected" in result.stdout
        assert "errors:" in result.stdout
        assert "errors: none" not in result.stdout


def test_error_report_next_cmd_points_to_strict() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        policy = root / "policy.yaml"
        out_dir = root / "out"

        _write_template(template, explicit_size=11)
        _write_task(task)
        _write_policy(policy, treat_inherited_as_error=False)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-mode",
                "report",
                "--format-baseline",
                "policy",
                "--policy",
                str(policy),
            ],
        )

        assert result.exit_code == 0
        assert "WARNING(format): format issues detected" in result.stdout
        assert "next_cmd:" in result.stdout
        assert "--preset strict" in result.stdout
