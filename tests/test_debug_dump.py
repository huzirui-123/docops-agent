from __future__ import annotations

from docx import Document

from core.render.debug_dump import collect_suspicious_runs


def test_collect_suspicious_runs_detects_symbol_font_and_touched_flag() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    symbol_run = paragraph.add_run("A(3F)")
    symbol_run.font.name = "Symbol"
    paragraph.add_run(" normal")

    report = collect_suspicious_runs(document, touched_runs={"p0:r0"}, stage="post_pipeline")

    assert report.scanned_runs == 2
    assert len(report.suspicious_runs) == 1
    run = report.suspicious_runs[0]
    assert run.run_id == "p0:r0"
    assert run.touched is True
    assert "symbol_font" in run.reasons
    assert run.font_info["name"] == "Symbol"
    assert run.codepoints[0].codepoint == "U+0041"


def test_collect_suspicious_runs_detects_unicode_category_cf() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("A\u202eB")

    report = collect_suspicious_runs(document, touched_runs=set(), stage="post_pipeline")

    assert len(report.suspicious_runs) == 1
    run = report.suspicious_runs[0]
    assert run.run_id == "p0:r0"
    assert "unicode_category:Cf" in run.reasons

