from __future__ import annotations

import json
from pathlib import Path

from docx.shared import Pt
from typer.testing import CliRunner

from apps.cli.main import app
from core.utils.docx_xml import set_run_fonts_and_size

runner = CliRunner()


def _write_template(path: Path) -> None:
    from docx import Document

    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("【MEETING_TITLE】")
    set_run_fonts_and_size(run, latin_font="Calibri", east_asia_font="宋体", size_pt=12)
    paragraph.paragraph_format.first_line_indent = Pt(24)  # 480 twips
    document.save(str(path))


def _write_task(path: Path) -> None:
    payload_json = json.dumps(
        {
            "task_type": "meeting_notice",
            "payload": {
                "meeting_title": "项目周会",
            },
        }
    )
    path.write_text(payload_json, encoding="utf-8")


def test_cli_fix_mode_none_keeps_indent_mismatch_and_returns_4() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out_none"
        _write_template(template)
        _write_task(task)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-mode",
                "strict",
                "--format-baseline",
                "policy",
                "--format-fix-mode",
                "none",
            ],
        )

        assert result.exit_code == 4
        report = json.loads((out_dir / "out.format_report.json").read_text(encoding="utf-8"))
        assert any(issue["code"] == "FIRST_LINE_INDENT_MISMATCH" for issue in report["issues"])
        assert report["summary"]["fix_applied"] is False


def test_cli_fix_mode_safe_removes_indent_mismatch() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out_safe"
        _write_template(template)
        _write_task(task)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-mode",
                "strict",
                "--format-baseline",
                "policy",
                "--format-fix-mode",
                "safe",
            ],
        )

        assert result.exit_code == 0
        report = json.loads((out_dir / "out.format_report.json").read_text(encoding="utf-8"))
        assert not any(issue["code"] == "FIRST_LINE_INDENT_MISMATCH" for issue in report["issues"])
        assert report["summary"]["fix_applied"] is True
        assert report["summary"]["fix_changes"]


def test_cli_off_mode_ignores_fix_mode() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out_off"
        _write_template(template)
        _write_task(task)

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--format-mode",
                "off",
                "--format-fix-mode",
                "safe",
            ],
        )

        assert result.exit_code == 0
        assert "INFO(format): fix skipped (off mode)" in result.stdout
        report = json.loads((out_dir / "out.format_report.json").read_text(encoding="utf-8"))
        assert report["summary"]["skipped"] is True
        assert report["summary"]["fix_applied"] is False

