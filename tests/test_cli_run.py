from __future__ import annotations

import json
from pathlib import Path

from typer.testing import CliRunner

from apps.cli.main import app

runner = CliRunner()


def _write_docx(path: Path, *, cross_run: bool = False, with_table: bool = False) -> None:
    from docx import Document

    document = Document()
    if with_table:
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).paragraphs[0].text = "table"
    else:
        paragraph = document.add_paragraph()
        paragraph.add_run("【MEETING_TITLE】")
        if cross_run:
            paragraph.add_run(" and ")
            paragraph.add_run("【BA")
            paragraph.add_run("D】")

    document.save(str(path))


def _write_task(path: Path, payload: dict[str, object]) -> None:
    payload_json = json.dumps({"task_type": "meeting_notice", "payload": payload})
    path.write_text(payload_json, encoding="utf-8")


def test_cli_success_writes_four_outputs(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    task = tmp_path / "task.json"
    out_dir = tmp_path / "out"
    _write_docx(template)
    _write_task(task, {"meeting_title": "Kickoff"})

    result = runner.invoke(
        app,
        [
            "run",
            "--template",
            str(template),
            "--task",
            str(task),
            "--skill",
            "meeting_notice",
            "--out-dir",
            str(out_dir),
        ],
    )

    assert result.exit_code == 0
    assert (out_dir / "out.docx").exists()
    assert (out_dir / "out.replace_log.json").exists()
    assert (out_dir / "out.missing_fields.json").exists()
    assert (out_dir / "out.format_report.json").exists()


def test_cli_template_error_returns_3() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out"
        _write_docx(template, cross_run=True)
        _write_task(task, {"meeting_title": "Kickoff"})

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
            ],
        )

        assert result.exit_code == 3
        replace_payload = json.loads((out_dir / "out.replace_log.json").read_text(encoding="utf-8"))
        assert replace_payload["summary"]["unsupported_count"] >= 1


def test_cli_missing_required_returns_2() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out"
        _write_docx(template)
        _write_task(task, {})

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
            ],
        )

        assert result.exit_code == 2
        missing_payload = json.loads(
            (out_dir / "out.missing_fields.json").read_text(encoding="utf-8")
        )
        assert missing_payload["missing_required"] == ["MEETING_TITLE"]


def test_cli_format_failed_returns_4() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out"
        _write_docx(template, with_table=True)
        _write_task(task, {})

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
            ],
        )

        assert result.exit_code == 4
        format_payload = json.loads(
            (out_dir / "out.format_report.json").read_text(encoding="utf-8")
        )
        assert format_payload["passed"] is False


def test_cli_exit_1_writes_error_fallback_json() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out"
        _write_docx(template)
        task.write_text("not-json", encoding="utf-8")

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
            ],
        )

        assert result.exit_code == 1
        replace_payload = json.loads((out_dir / "out.replace_log.json").read_text(encoding="utf-8"))
        assert "error" in replace_payload
        assert replace_payload["error"]["error_type"]


def test_cli_no_overwrite_returns_1() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out"
        out_dir.mkdir()
        _write_docx(template)
        _write_task(task, {"meeting_title": "Kickoff"})
        existing_docx = b"existing-docx-content"
        existing_replace = '{"marker":"do-not-overwrite"}'
        (out_dir / "out.docx").write_bytes(existing_docx)
        (out_dir / "out.replace_log.json").write_text(existing_replace, encoding="utf-8")

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--no-overwrite",
            ],
        )

        assert result.exit_code == 1
        assert "--no-overwrite" in result.stdout
        assert (out_dir / "out.docx").read_bytes() == existing_docx
        assert (out_dir / "out.replace_log.json").read_text(encoding="utf-8") == existing_replace
        assert not (out_dir / "out.missing_fields.json").exists()
        assert not (out_dir / "out.format_report.json").exists()


def test_cli_default_overwrite_prints_info() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out"
        out_dir.mkdir()
        _write_docx(template)
        _write_task(task, {"meeting_title": "Kickoff"})
        (out_dir / "out.docx").write_bytes(b"exists")

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
            ],
        )

        assert result.exit_code == 0
        assert "INFO: overwriting existing outputs" in result.stdout


def test_cli_warn_mode_prints_warning() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out"
        _write_docx(template, cross_run=True)
        _write_task(task, {"meeting_title": "Kickoff"})

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--unsupported-mode",
                "warn",
            ],
        )

        assert result.exit_code == 0
        assert "WARNING: unsupported placeholders detected" in result.stdout
