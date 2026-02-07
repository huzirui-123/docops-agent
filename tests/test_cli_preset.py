from __future__ import annotations

import json
from pathlib import Path

from typer.testing import CliRunner

from apps.cli.main import app

runner = CliRunner()


def _write_task(path: Path, payload: dict[str, object]) -> None:
    payload_json = json.dumps({"task_type": "meeting_notice", "payload": payload})
    path.write_text(payload_json, encoding="utf-8")


def _write_simple_template(path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("【MEETING_TITLE】")
    document.save(str(path))


def _write_table_template(path: Path) -> None:
    from docx import Document

    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].text = "table-body"
    document.save(str(path))


def test_cli_default_and_quick_preset_show_expected_effective_modes() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_default = root / "out_default"
        out_quick = root / "out_quick"
        _write_simple_template(template)
        _write_task(task, {"meeting_title": "Kickoff"})

        default_result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_default),
            ],
        )

        quick_result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_quick),
                "--preset",
                "quick",
            ],
        )

        assert default_result.exit_code == 0
        assert "format_mode=report" in default_result.stdout
        assert "format_baseline=template" in default_result.stdout
        assert "format_fix_mode=safe" in default_result.stdout

        assert quick_result.exit_code == 0
        assert "format_mode=report" in quick_result.stdout
        assert "format_baseline=template" in quick_result.stdout
        assert "format_fix_mode=safe" in quick_result.stdout


def test_cli_template_preset_sets_fix_mode_none() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out_template"
        _write_simple_template(template)
        _write_task(task, {"meeting_title": "Kickoff"})

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--preset",
                "template",
            ],
        )

        assert result.exit_code == 0
        assert "format_fix_mode=none" in result.stdout


def test_cli_strict_preset_blocks_on_format_issues() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out_strict"
        _write_table_template(template)
        _write_task(task, {})

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--preset",
                "strict",
            ],
        )

        assert result.exit_code == 4


def test_cli_preset_conflict_with_advanced_flags_returns_1_and_writes_nothing() -> None:
    with runner.isolated_filesystem() as tmp:
        root = Path(tmp)
        template = root / "template.docx"
        task = root / "task.json"
        out_dir = root / "out_conflict"
        _write_simple_template(template)
        _write_task(task, {"meeting_title": "Kickoff"})

        result = runner.invoke(
            app,
            [
                "run",
                "--template",
                str(template),
                "--task",
                str(task),
                "--skill",
                "meeting_notice",
                "--out-dir",
                str(out_dir),
                "--preset",
                "strict",
                "--format-mode",
                "report",
            ],
        )

        assert result.exit_code == 1
        assert "--preset cannot be combined" in result.stdout
        assert not (out_dir / "out.docx").exists()
        assert not (out_dir / "out.replace_log.json").exists()
        assert not (out_dir / "out.missing_fields.json").exists()
        assert not (out_dir / "out.format_report.json").exists()

