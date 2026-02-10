from __future__ import annotations

import io
import json
import zipfile

import httpx
import pytest
from docx import Document

from apps.api.main import app


def _build_docx_bytes(*, with_table: bool = False, with_placeholder: bool = True) -> bytes:
    document = Document()
    if with_table:
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).paragraphs[0].text = "table-content"
    elif with_placeholder:
        document.add_paragraph("【MEETING_TITLE】")
    else:
        document.add_paragraph("plain")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _task_bytes() -> bytes:
    payload = {"task_type": "meeting_notice", "payload": {"meeting_title": "Kickoff"}}
    return json.dumps(payload).encode("utf-8")


def _zip_map(content: bytes) -> dict[str, bytes]:
    with zipfile.ZipFile(io.BytesIO(content), "r") as archive:
        return {name: archive.read(name) for name in archive.namelist()}


def _assert_timing_payload(payload: dict[str, object]) -> None:
    timing = payload["timing"]
    assert isinstance(timing, dict)
    for key in ("total_ms", "subprocess_ms", "zip_ms", "queue_wait_ms"):
        value = timing[key]
        assert isinstance(value, int)
        assert value >= 0


@pytest.mark.anyio
async def test_healthz_ok() -> None:
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
        response = await client.get("/healthz")

    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


@pytest.mark.anyio
async def test_run_quick_returns_zip_with_required_files() -> None:
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
        response = await client.post(
            "/v1/run",
            files={
                "template": (
                    "template.docx",
                    _build_docx_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                "task": ("task.json", _task_bytes(), "application/json"),
            },
            data={"skill": "meeting_notice"},
        )

    assert response.status_code == 200
    assert response.headers["X-Docops-Exit-Code"] == "0"
    request_id = response.headers["X-Docops-Request-Id"]
    assert request_id

    zipped = _zip_map(response.content)
    required = {
        "out.docx",
        "out.replace_log.json",
        "out.missing_fields.json",
        "out.format_report.json",
        "api_result.json",
    }
    assert required.issubset(set(zipped.keys()))

    api_result = json.loads(zipped["api_result.json"].decode("utf-8"))
    assert api_result["exit_code"] == 0
    assert api_result["effective"]["format_mode"] == "report"
    assert api_result["effective"]["format_report"] == "json"
    assert api_result["request_id"] == request_id
    _assert_timing_payload(api_result)


@pytest.mark.anyio
async def test_run_strict_format_fail_returns_zip_with_exit_code_4() -> None:
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
        response = await client.post(
            "/v1/run",
            files={
                "template": (
                    "template.docx",
                    _build_docx_bytes(with_table=True, with_placeholder=False),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                "task": ("task.json", _task_bytes(), "application/json"),
            },
            data={"skill": "meeting_notice", "preset": "strict"},
        )

    assert response.status_code == 200
    assert response.headers["X-Docops-Exit-Code"] == "4"
    request_id = response.headers["X-Docops-Request-Id"]

    zipped = _zip_map(response.content)
    assert "out.docx" in zipped
    assert "out.format_report.json" in zipped
    assert "api_result.json" in zipped

    api_result = json.loads(zipped["api_result.json"].decode("utf-8"))
    assert api_result["exit_code"] == 4
    assert api_result["effective"]["format_mode"] == "strict"
    assert api_result["request_id"] == request_id
    _assert_timing_payload(api_result)


@pytest.mark.anyio
async def test_run_conflict_returns_400_json() -> None:
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
        response = await client.post(
            "/v1/run",
            files={
                "template": (
                    "template.docx",
                    _build_docx_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                "task": ("task.json", _task_bytes(), "application/json"),
            },
            data={
                "skill": "meeting_notice",
                "preset": "strict",
                "format_mode": "report",
            },
        )

    assert response.status_code == 400
    assert response.headers["X-Docops-Request-Id"]
    payload = response.json()
    assert payload["error_code"] == "INVALID_ARGUMENT_CONFLICT"
    assert payload["detail"]["request_id"] == response.headers["X-Docops-Request-Id"]


@pytest.mark.anyio
async def test_run_invalid_docx_returns_415() -> None:
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
        response = await client.post(
            "/v1/run",
            files={
                "template": ("template.docx", b"not-a-docx", "application/octet-stream"),
                "task": ("task.json", _task_bytes(), "application/json"),
            },
            data={"skill": "meeting_notice"},
        )

    assert response.status_code == 415
    assert response.headers["X-Docops-Request-Id"]
    payload = response.json()
    assert payload["error_code"] == "INVALID_MEDIA_TYPE"
    assert payload["detail"]["request_id"] == response.headers["X-Docops-Request-Id"]


@pytest.mark.anyio
async def test_run_invalid_task_json_returns_400() -> None:
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
        response = await client.post(
            "/v1/run",
            files={
                "template": (
                    "template.docx",
                    _build_docx_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                "task": ("task.json", b"not-json", "application/json"),
            },
            data={"skill": "meeting_notice"},
        )

    assert response.status_code == 400
    assert response.headers["X-Docops-Request-Id"]
    payload = response.json()
    assert payload["error_code"] == "INVALID_JSON"
    assert payload["detail"]["request_id"] == response.headers["X-Docops-Request-Id"]


@pytest.mark.anyio
async def test_zip_contains_fixed_set_and_api_result_matches_header() -> None:
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
        response = await client.post(
            "/v1/run",
            files={
                "template": (
                    "template.docx",
                    _build_docx_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                "task": ("task.json", _task_bytes(), "application/json"),
            },
            data={"skill": "meeting_notice", "export_suggested_policy": "true"},
        )

    assert response.status_code == 200
    request_id = response.headers["X-Docops-Request-Id"]

    zipped = _zip_map(response.content)
    names = set(zipped.keys())
    required = {
        "out.docx",
        "out.replace_log.json",
        "out.missing_fields.json",
        "out.format_report.json",
        "api_result.json",
    }
    assert required.issubset(names)
    assert "out.suggested_policy.yaml" in names

    api_result = json.loads(zipped["api_result.json"].decode("utf-8"))
    assert str(api_result["exit_code"]) == response.headers["X-Docops-Exit-Code"]
    assert api_result["request_id"] == request_id
    _assert_timing_payload(api_result)


@pytest.mark.anyio
async def test_debug_artifacts_include_trace(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("DOCOPS_DEBUG_ARTIFACTS", "1")

    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
        response = await client.post(
            "/v1/run",
            files={
                "template": (
                    "template.docx",
                    _build_docx_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                "task": ("task.json", _task_bytes(), "application/json"),
            },
            data={"skill": "meeting_notice"},
        )

    assert response.status_code == 200
    request_id = response.headers["X-Docops-Request-Id"]

    zipped = _zip_map(response.content)
    assert "trace.json" in zipped

    trace_payload = json.loads(zipped["trace.json"].decode("utf-8"))
    assert trace_payload["request_id"] == request_id
    assert isinstance(trace_payload["timing"], dict)
