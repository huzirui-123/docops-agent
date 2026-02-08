"""Format validator using policy constraints and touched-run scope."""

from __future__ import annotations

from collections.abc import Mapping
from typing import Literal

from docx.document import Document as DocxDocument

from core.format.models import FormatIssue, FormatPolicy, FormatReport
from core.format.run_ids import ParagraphRunContext, iter_paragraph_run_contexts
from core.utils.docx_xml import (
    get_first_line_indent_twips,
    get_line_spacing_twips,
    get_run_east_asia_font,
    paragraph_has_direct_numpr,
)


def validate_document(
    document: DocxDocument,
    policy: FormatPolicy,
    touched_runs: set[str],
    *,
    baseline: Literal["template", "policy"] = "policy",
    template_run_styles: Mapping[str, object] | None = None,
) -> FormatReport:
    """Validate document format with policy.

    Notes:
    - When forbid_tables is true, table presence is reported as non-fixable and
      table cell paragraphs are skipped for all other checks.
    - numPr checks only inspect direct pPr.numPr, not style inheritance.
    - When first_line_indent_twips is None, first-line indent validation is skipped.
    - RUN font/size checks are baseline-aware and observability-aware.
    """

    issues: list[FormatIssue] = []
    run_style_snapshots = template_run_styles or {}

    if policy.forbid_tables and document.tables:
        table_count = len(document.tables)
        for table_index, _ in enumerate(document.tables):
            issues.append(
                FormatIssue(
                    code="TABLE_FORBIDDEN",
                    message="Tables are forbidden by policy.",
                    paragraph_path=f"t{table_index}",
                    fixable=False,
                    context={"table_count": table_count},
                )
            )

    include_tables = not policy.forbid_tables
    for context in iter_paragraph_run_contexts(document, include_tables=include_tables):
        issues.extend(
            _validate_paragraph(
                context,
                policy,
                touched_runs,
                baseline=baseline,
                template_run_styles=run_style_snapshots,
            )
        )

    return FormatReport.from_issues(issues)


def _validate_paragraph(
    context: ParagraphRunContext,
    policy: FormatPolicy,
    touched_runs: set[str],
    *,
    baseline: Literal["template", "policy"],
    template_run_styles: Mapping[str, object],
) -> list[FormatIssue]:
    issues: list[FormatIssue] = []
    paragraph = context.paragraph

    if policy.forbid_numpr and paragraph_has_direct_numpr(paragraph):
        issues.append(
            FormatIssue(
                code="NUMPR_PRESENT",
                message="Direct paragraph numbering (pPr.numPr) is present.",
                paragraph_path=context.paragraph_path,
                fixable=True,
            )
        )

    actual_line_spacing = get_line_spacing_twips(paragraph)
    if actual_line_spacing is not None and not _within_tolerance(
        actual_line_spacing, policy.line_spacing_twips, policy.twips_tolerance
    ):
        issues.append(
            FormatIssue(
                code="LINE_SPACING_MISMATCH",
                message=(
                    f"Line spacing twips mismatch: expected {policy.line_spacing_twips}, "
                    f"got {actual_line_spacing}."
                ),
                paragraph_path=context.paragraph_path,
                fixable=False,
            )
        )

    if policy.first_line_indent_twips is not None:
        actual_indent = get_first_line_indent_twips(paragraph)
        if actual_indent is not None and not _within_tolerance(
            actual_indent, policy.first_line_indent_twips, policy.twips_tolerance
        ):
            issues.append(
                FormatIssue(
                    code="FIRST_LINE_INDENT_MISMATCH",
                    message=(
                        "First-line indent twips mismatch: "
                        f"expected {policy.first_line_indent_twips}, got {actual_indent}."
                    ),
                    paragraph_path=context.paragraph_path,
                    fixable=False,
                    expected=policy.first_line_indent_twips,
                    actual=actual_indent,
                    tolerance=policy.twips_tolerance,
                )
            )

    if policy.trim_leading_spaces and _has_leading_trim_chars(paragraph, set(policy.trim_chars)):
        issues.append(
            FormatIssue(
                code="LEADING_WHITESPACE",
                message="Paragraph has trim-target leading whitespace.",
                paragraph_path=context.paragraph_path,
                fixable=True,
            )
        )

    for run_id, run in zip(context.run_ids, paragraph.runs, strict=False):
        if run_id not in touched_runs:
            continue

        latin_font = run.font.name
        east_asia_font = get_run_east_asia_font(run)
        size = run.font.size.pt if run.font.size is not None else None
        size_pt = None if size is None else int(round(size))
        template_style = template_run_styles.get(run_id)

        font_issue = _build_font_issue(
            paragraph_path=context.paragraph_path,
            run_id=run_id,
            baseline=baseline,
            policy=policy,
            template_style=template_style,
            latin_font=latin_font,
            east_asia_font=east_asia_font,
        )
        if font_issue is not None:
            issues.append(font_issue)

        size_issue = _build_size_issue(
            paragraph_path=context.paragraph_path,
            run_id=run_id,
            baseline=baseline,
            policy=policy,
            template_style=template_style,
            rendered_size_pt=size_pt,
        )
        if size_issue is not None:
            issues.append(size_issue)

    return issues


def _has_leading_trim_chars(paragraph, trim_chars: set[str]) -> bool:
    if not paragraph.runs:
        return False

    prefix = ""
    for run in paragraph.runs:
        prefix += run.text or ""
        if prefix:
            break

    return bool(prefix) and prefix[0] in trim_chars


def _within_tolerance(actual: int, expected: int, tolerance: int) -> bool:
    return abs(actual - expected) <= tolerance


def _build_font_issue(
    *,
    paragraph_path: str,
    run_id: str,
    baseline: Literal["template", "policy"],
    policy: FormatPolicy,
    template_style: object | None,
    latin_font: str | None,
    east_asia_font: str | None,
) -> FormatIssue | None:
    if baseline == "template":
        expected_latin = _snapshot_get(template_style, "latin_font")
        expected_east_asia = _snapshot_get(template_style, "east_asia_font")
        mismatch, severity, observability = _compare_font_against_template(
            expected_latin=expected_latin,
            expected_east_asia=expected_east_asia,
            actual_latin=latin_font,
            actual_east_asia=east_asia_font,
        )
        if not mismatch:
            return None
        return FormatIssue(
            code="RUN_FONT_MISMATCH",
            message=(
                "Touched run font mismatch against template snapshot: "
                f"expected latin={expected_latin!r}, eastAsia={expected_east_asia!r}; "
                f"got latin={latin_font!r}, eastAsia={east_asia_font!r}."
            ),
            paragraph_path=paragraph_path,
            run_id=run_id,
            fixable=True,
            expected={"latin": expected_latin, "east_asia": expected_east_asia},
            actual={"latin": latin_font, "east_asia": east_asia_font},
            template_value={"latin": expected_latin, "east_asia": expected_east_asia},
            rendered_value={"latin": latin_font, "east_asia": east_asia_font},
            context={"baseline": "template"},
            severity=severity,
            observability=observability,
        )

    mismatch, severity, observability = _compare_font_against_policy(
        expected_latin=policy.run_font_latin,
        expected_east_asia=policy.run_font_east_asia,
        actual_latin=latin_font,
        actual_east_asia=east_asia_font,
        treat_inherited_as_error=policy.treat_inherited_as_error,
    )
    if not mismatch:
        return None
    return FormatIssue(
        code="RUN_FONT_MISMATCH",
        message=(
            "Touched run font mismatch against policy: "
            f"expected latin={policy.run_font_latin!r}, eastAsia={policy.run_font_east_asia!r}; "
            f"got latin={latin_font!r}, eastAsia={east_asia_font!r}."
        ),
        paragraph_path=paragraph_path,
        run_id=run_id,
        fixable=True,
        expected={"latin": policy.run_font_latin, "east_asia": policy.run_font_east_asia},
        actual={"latin": latin_font, "east_asia": east_asia_font},
        rendered_value={"latin": latin_font, "east_asia": east_asia_font},
        context={"baseline": "policy"},
        severity=severity,
        observability=observability,
    )


def _build_size_issue(
    *,
    paragraph_path: str,
    run_id: str,
    baseline: Literal["template", "policy"],
    policy: FormatPolicy,
    template_style: object | None,
    rendered_size_pt: int | None,
) -> FormatIssue | None:
    if baseline == "template":
        expected_size_pt = _snapshot_get_int(template_style, "size_pt")
        mismatch, severity, observability = _compare_scalar_against_template(
            expected=expected_size_pt,
            actual=rendered_size_pt,
        )
        if not mismatch:
            return None
        return FormatIssue(
            code="RUN_SIZE_MISMATCH",
            message=(
                "Touched run size mismatch against template snapshot: "
                f"expected {expected_size_pt}pt, got {rendered_size_pt}pt."
            ),
            paragraph_path=paragraph_path,
            run_id=run_id,
            fixable=True,
            expected=expected_size_pt,
            actual=rendered_size_pt,
            template_value=expected_size_pt,
            rendered_value=rendered_size_pt,
            context={"baseline": "template"},
            severity=severity,
            observability=observability,
        )

    mismatch, severity, observability = _compare_scalar_against_policy(
        expected=policy.run_size_pt,
        actual=rendered_size_pt,
        treat_inherited_as_error=policy.treat_inherited_as_error,
    )
    if not mismatch:
        return None
    return FormatIssue(
        code="RUN_SIZE_MISMATCH",
        message=(
            "Touched run size mismatch against policy: "
            f"expected {policy.run_size_pt}pt, got {rendered_size_pt}pt."
        ),
        paragraph_path=paragraph_path,
        run_id=run_id,
        fixable=True,
        expected=policy.run_size_pt,
        actual=rendered_size_pt,
        rendered_value=rendered_size_pt,
        context={"baseline": "policy"},
        severity=severity,
        observability=observability,
    )


def _compare_font_against_template(
    *,
    expected_latin: str | None,
    expected_east_asia: str | None,
    actual_latin: str | None,
    actual_east_asia: str | None,
) -> tuple[bool, Literal["error", "warn"], Literal["observed", "unknown"]]:
    error = False
    warn = False

    for expected, actual in (
        (expected_latin, actual_latin),
        (expected_east_asia, actual_east_asia),
    ):
        if expected is None and actual is None:
            continue
        if expected is None and actual is not None:
            error = True
            continue
        if expected is not None and actual is None:
            warn = True
            continue
        if expected != actual:
            error = True

    if error:
        return True, "error", "observed"
    if warn:
        return True, "warn", "unknown"
    return False, "error", "observed"


def _compare_font_against_policy(
    *,
    expected_latin: str,
    expected_east_asia: str,
    actual_latin: str | None,
    actual_east_asia: str | None,
    treat_inherited_as_error: bool,
) -> tuple[bool, Literal["error", "warn"], Literal["observed", "unknown"]]:
    has_error = False
    has_warn = False
    unknown = False

    for expected, actual in (
        (expected_latin, actual_latin),
        (expected_east_asia, actual_east_asia),
    ):
        if actual is None:
            if treat_inherited_as_error:
                has_error = True
            else:
                has_warn = True
            unknown = True
            continue
        if actual != expected:
            has_error = True

    if has_error:
        if unknown and not has_warn:
            return True, "error", "unknown"
        return True, "error", "observed"
    if has_warn:
        return True, "warn", "unknown"
    return False, "error", "observed"


def _compare_scalar_against_template(
    *,
    expected: int | None,
    actual: int | None,
) -> tuple[bool, Literal["error", "warn"], Literal["observed", "unknown"]]:
    if expected is None and actual is None:
        return False, "error", "observed"
    if expected is None and actual is not None:
        return True, "error", "observed"
    if expected is not None and actual is None:
        return True, "warn", "unknown"
    if expected != actual:
        return True, "error", "observed"
    return False, "error", "observed"


def _compare_scalar_against_policy(
    *,
    expected: int,
    actual: int | None,
    treat_inherited_as_error: bool,
) -> tuple[bool, Literal["error", "warn"], Literal["observed", "unknown"]]:
    if actual is None:
        if treat_inherited_as_error:
            return True, "error", "unknown"
        return True, "warn", "unknown"
    if expected != actual:
        return True, "error", "observed"
    return False, "error", "observed"


def _snapshot_get(snapshot: object | None, field: str) -> str | None:
    if snapshot is None:
        return None
    if isinstance(snapshot, dict):
        value = snapshot.get(field)
        return value if isinstance(value, str) else None
    value = getattr(snapshot, field, None)
    return value if isinstance(value, str) else None


def _snapshot_get_int(snapshot: object | None, field: str) -> int | None:
    if snapshot is None:
        return None
    if isinstance(snapshot, dict):
        value = snapshot.get(field)
    else:
        value = getattr(snapshot, field, None)
    if isinstance(value, int):
        return value
    return None
