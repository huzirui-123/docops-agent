"""Observed formatting snapshots for template-vs-rendered reporting."""

from __future__ import annotations

from collections.abc import Iterable
from typing import cast

from docx.document import Document as DocxDocument

from core.format.models import FormatObserved, FormatObservedDiff
from core.format.run_ids import iter_paragraph_run_contexts
from core.utils.docx_xml import get_run_font_info, paragraph_has_direct_numpr


def observe_document(document: DocxDocument) -> FormatObserved:
    """Collect high-level formatting observations from body and table content."""

    indent_hist: dict[str, int] = {}
    latin_hist: dict[str, int] = {}
    east_hist: dict[str, int] = {}
    has_numpr = False

    for context in iter_paragraph_run_contexts(document, include_tables=True):
        if paragraph_has_direct_numpr(context.paragraph):
            has_numpr = True

        indent_key = _indent_to_key(context.paragraph.paragraph_format.first_line_indent)
        _inc(indent_hist, indent_key)

        for run in context.paragraph.runs:
            latin_font = run.font.name if run.font.name is not None else "none"
            _inc(latin_hist, latin_font)

            font_info = get_run_font_info(run)
            east_asia = font_info.get("eastAsia")
            _inc(east_hist, east_asia if east_asia is not None else "unknown")

    return FormatObserved(
        has_tables=bool(document.tables),
        has_numpr=has_numpr,
        first_line_indent_twips_hist=indent_hist,
        run_font_latin_hist=latin_hist,
        run_font_east_asia_hist=east_hist,
    )


def dominant_first_line_indent_twips(document: DocxDocument) -> int | None:
    """Pick dominant first-line indent with body-first fallback.

    Strategy:
    - Prefer body paragraphs (`document.paragraphs`).
    - Ignore "none" and tiny values (< 200 twips).
    - If body has no usable indent values, fall back to table paragraphs.
    """

    body_hist = _collect_indent_hist(document.paragraphs)
    dominant = pick_dominant_indent_from_hist(body_hist)
    if dominant is not None:
        return dominant

    table_paragraphs = (
        paragraph
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    table_hist = _collect_indent_hist(table_paragraphs)
    return pick_dominant_indent_from_hist(table_hist)


def diff_observed(template: FormatObserved, rendered: FormatObserved) -> FormatObservedDiff:
    """Compute rendered-template delta from observed snapshots."""

    keys = set(template.first_line_indent_twips_hist) | set(rendered.first_line_indent_twips_hist)
    indent_delta: dict[str, int] = {}
    for key in sorted(keys):
        indent_delta[key] = rendered.first_line_indent_twips_hist.get(
            key, 0
        ) - template.first_line_indent_twips_hist.get(key, 0)

    return FormatObservedDiff(
        has_tables_changed=template.has_tables != rendered.has_tables,
        has_numpr_changed=template.has_numpr != rendered.has_numpr,
        first_line_indent_twips_hist_delta=indent_delta,
    )


def pick_dominant_indent_from_hist(indent_hist: dict[str, int], min_twips: int = 200) -> int | None:
    """Pick dominant indent value from a histogram.

    Keys must be twips as strings; "none" and values below min_twips are ignored.
    """

    numeric_entries: list[tuple[int, int]] = []
    for key, count in indent_hist.items():
        if key == "none":
            continue
        try:
            twips = int(key)
        except ValueError:
            continue
        if twips < min_twips:
            continue
        numeric_entries.append((twips, count))

    if not numeric_entries:
        return None

    numeric_entries.sort(key=lambda item: (-item[1], item[0]))
    return numeric_entries[0][0]


def _indent_to_key(indent) -> str:
    if indent is None:
        return "none"

    twips = getattr(indent, "twips", None)
    if isinstance(twips, int):
        return str(twips)

    pt = getattr(indent, "pt", None)
    if isinstance(pt, int | float):
        return str(int(round(cast(float, pt) * 20)))

    return "none"


def _inc(hist: dict[str, int], key: str) -> None:
    hist[key] = hist.get(key, 0) + 1


def _collect_indent_hist(paragraphs: Iterable) -> dict[str, int]:
    hist: dict[str, int] = {}
    for paragraph in paragraphs:
        key = _indent_to_key(paragraph.paragraph_format.first_line_indent)
        _inc(hist, key)
    return hist
