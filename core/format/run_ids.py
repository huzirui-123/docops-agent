"""Run ID helpers that keep format tooling aligned with M2 run_id convention."""

from __future__ import annotations

from collections.abc import Iterator
from dataclasses import dataclass

from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph


@dataclass(frozen=True)
class ParagraphRunContext:
    """Paragraph metadata and run IDs using M2 convention."""

    paragraph: Paragraph
    paragraph_path: str
    run_ids: list[str]
    in_table: bool


def iter_paragraph_run_contexts(
    document: DocxDocument, *, include_tables: bool = True
) -> Iterator[ParagraphRunContext]:
    """Yield body/table paragraph contexts with stable run IDs.

    Run ID format matches M2 parser exactly:
    - body: p{p_idx}:r{r_idx}
    - table: t{t_idx}.r{row_idx}.c{cell_idx}.p{p_idx}:r{r_idx}
    """

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        paragraph_path = f"p{paragraph_index}"
        run_ids = [
            f"{paragraph_path}:r{run_index}" for run_index, _ in enumerate(paragraph.runs)
        ]
        yield ParagraphRunContext(
            paragraph=paragraph,
            paragraph_path=paragraph_path,
            run_ids=run_ids,
            in_table=False,
        )

    if not include_tables:
        return

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    paragraph_path = (
                        f"t{table_index}.r{row_index}.c{cell_index}.p{paragraph_index}"
                    )
                    run_ids = [
                        f"{paragraph_path}:r{run_index}"
                        for run_index, _ in enumerate(paragraph.runs)
                    ]
                    yield ParagraphRunContext(
                        paragraph=paragraph,
                        paragraph_path=paragraph_path,
                        run_ids=run_ids,
                        in_table=True,
                    )
