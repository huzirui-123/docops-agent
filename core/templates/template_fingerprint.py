"""Template fingerprint generation that is insensitive to run segmentation."""

from __future__ import annotations

import hashlib
import json
import re
from dataclasses import asdict
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument

from core.templates.models import (
    FingerprintOccurrence,
    FingerprintParagraph,
    FingerprintPayload,
    FingerprintUnsupported,
    Occurrence,
    UnsupportedOccurrence,
)
from core.templates.placeholder_parser import parse_placeholders

_SPACE_RE = re.compile(r"[ \t]+")


def compute_template_fingerprint(document: DocxDocument) -> str:
    """Compute a canonical SHA256 fingerprint for a template document."""

    paragraph_texts, run_offsets = _collect_paragraph_entries(document)
    parse_result = parse_placeholders(document, strict=False)

    payload = FingerprintPayload(
        paragraphs=sorted(
            [
                FingerprintParagraph(
                    paragraph_path=path,
                    text=_normalize_whitespace(text),
                )
                for path, text in paragraph_texts.items()
            ],
            key=lambda item: item.paragraph_path,
        ),
        occurrences=_build_occurrences(parse_result.occurrences, run_offsets),
        unsupported=_build_unsupported(parse_result.unsupported),
    )

    serialized = json.dumps(asdict(payload), sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(serialized.encode("utf-8")).hexdigest()


def compute_template_fingerprint_from_path(path: Path) -> str:
    """Load a .docx file and compute its fingerprint."""

    document = Document(str(path))
    return compute_template_fingerprint(document)


def _collect_paragraph_entries(
    document: DocxDocument,
) -> tuple[dict[str, str], dict[str, list[int]]]:
    paragraph_texts: dict[str, str] = {}
    run_offsets: dict[str, list[int]] = {}

    for path, paragraph in _iter_target_paragraphs(document):
        full_text = "".join((run.text or "") for run in paragraph.runs)
        offsets: list[int] = []
        cursor = 0
        for run in paragraph.runs:
            offsets.append(cursor)
            cursor += len(run.text or "")

        paragraph_texts[path] = full_text
        run_offsets[path] = offsets

    return paragraph_texts, run_offsets


def _iter_target_paragraphs(document: DocxDocument):
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        yield f"p{paragraph_index}", paragraph

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    paragraph_path = (
                        f"t{table_index}.r{row_index}.c{cell_index}.p{paragraph_index}"
                    )
                    yield paragraph_path, paragraph


def _build_occurrences(
    occurrences: list[Occurrence], run_offsets: dict[str, list[int]]
) -> list[FingerprintOccurrence]:
    payload_occurrences: list[FingerprintOccurrence] = []

    for occurrence in occurrences:
        paragraph_path, run_index = _parse_run_id(occurrence.run_id)
        offsets = run_offsets[paragraph_path]
        run_start = offsets[run_index]
        payload_occurrences.append(
            FingerprintOccurrence(
                paragraph_path=paragraph_path,
                start=run_start + occurrence.start,
                end=run_start + occurrence.end,
                field_name=occurrence.field_name,
            )
        )

    return sorted(
        payload_occurrences,
        key=lambda item: (item.paragraph_path, item.start, item.end, item.field_name),
    )


def _build_unsupported(
    unsupported_items: list[UnsupportedOccurrence],
) -> list[FingerprintUnsupported]:
    payload_unsupported: list[FingerprintUnsupported] = []

    for item in unsupported_items:
        paragraph_path = _paragraph_path_from_run_id(item.run_id)
        payload_unsupported.append(
            FingerprintUnsupported(
                kind=item.kind,
                text=_normalize_whitespace(item.text),
                paragraph_path=paragraph_path,
                start=item.start,
                end=item.end,
            )
        )

    return sorted(
        payload_unsupported,
        key=lambda item: (
            item.paragraph_path or "",
            item.start if item.start is not None else -1,
            item.end if item.end is not None else -1,
            item.kind,
            item.text,
        ),
    )


def _parse_run_id(run_id: str) -> tuple[str, int]:
    paragraph_path, run_suffix = run_id.rsplit(":r", maxsplit=1)
    return paragraph_path, int(run_suffix)


def _paragraph_path_from_run_id(run_id: str | None) -> str | None:
    if run_id is None:
        return None
    paragraph_path, _ = run_id.rsplit(":r", maxsplit=1)
    return paragraph_path


def _normalize_whitespace(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _SPACE_RE.sub(" ", text)
    return text.strip()
