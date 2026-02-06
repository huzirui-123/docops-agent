"""Placeholder parser constrained to body paragraphs and table cells.

Header and footer content is intentionally ignored in M2.
"""

from __future__ import annotations

import re
from collections.abc import Iterator

from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from core.templates.models import Occurrence, ParseResult, UnsupportedOccurrence
from core.utils.errors import TemplateError

_VALID_PLACEHOLDER_RE = re.compile(r"【([A-Z0-9_]+)】")
_BRACKETED_RE = re.compile(r"【([^】]*)】")
_FIELD_NAME_RE = re.compile(r"[A-Z0-9_]+")
_OPEN_BRACKET = "【"
_CLOSE_BRACKET = "】"


def parse_placeholders(document: DocxDocument, strict: bool = False) -> ParseResult:
    """Parse placeholders from document body and tables.

    Rules:
    - Supported placeholder format is exactly: 【FIELD_NAME】
    - FIELD_NAME allows only A-Z, 0-9 and underscore.
    - Placeholders crossing runs are marked unsupported.
    - Header/footer are ignored in this stage.

    Args:
        document: python-docx document object.
        strict: When True, raise TemplateError if any unsupported item exists.

    Returns:
        ParseResult containing fields, supported occurrences and unsupported entries.
    """

    result = ParseResult()
    seen_fields: set[str] = set()

    for paragraph, run_id_pattern in _iter_target_paragraphs(document):
        full_text, run_spans = _build_run_spans(paragraph)
        if not full_text:
            continue

        for match in _VALID_PLACEHOLDER_RE.finditer(full_text):
            field_name = match.group(1)
            start_run = _run_index_for_position(match.start(), run_spans)
            end_run = _run_index_for_position(match.end() - 1, run_spans)
            if start_run is None or end_run is None or start_run != end_run:
                result.unsupported.append(
                    UnsupportedOccurrence(
                        kind="cross_run",
                        text=match.group(0),
                        run_id=run_id_pattern.format(start_run) if start_run is not None else None,
                        start=match.start(),
                        end=match.end(),
                    )
                )
                continue

            run_start = run_spans[start_run][1]
            occurrence = Occurrence(
                field_name=field_name,
                run_id=run_id_pattern.format(start_run),
                start=match.start() - run_start,
                end=match.end() - run_start,
            )
            result.occurrences.append(occurrence)
            if field_name not in seen_fields:
                result.fields.append(field_name)
                seen_fields.add(field_name)

        for match in _BRACKETED_RE.finditer(full_text):
            token = match.group(0)
            inner = match.group(1)
            if _FIELD_NAME_RE.fullmatch(inner):
                continue
            start_run = _run_index_for_position(match.start(), run_spans)
            result.unsupported.append(
                UnsupportedOccurrence(
                    kind="invalid_format",
                    text=token,
                    run_id=run_id_pattern.format(start_run) if start_run is not None else None,
                    start=match.start(),
                    end=match.end(),
                )
            )

        for kind, start, end, text in _find_unbalanced_brackets(full_text):
            run_index = _run_index_for_position(start, run_spans)
            result.unsupported.append(
                UnsupportedOccurrence(
                    kind=kind,
                    text=text,
                    run_id=run_id_pattern.format(run_index) if run_index is not None else None,
                    start=start,
                    end=end,
                )
            )

    if strict and result.unsupported:
        raise TemplateError("Unsupported placeholders found in template", result=result)

    return result


def _iter_target_paragraphs(document: DocxDocument) -> Iterator[tuple[Paragraph, str]]:
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        yield paragraph, f"p{paragraph_index}:r{{}}"

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    run_id_pattern = (
                        f"t{table_index}.r{row_index}.c{cell_index}.p{paragraph_index}:r{{}}"
                    )
                    yield paragraph, run_id_pattern


def _build_run_spans(paragraph: Paragraph) -> tuple[str, list[tuple[int, int, int]]]:
    run_spans: list[tuple[int, int, int]] = []
    chunks: list[str] = []
    cursor = 0

    for run_index, run in enumerate(paragraph.runs):
        text = run.text or ""
        start = cursor
        cursor += len(text)
        run_spans.append((run_index, start, cursor))
        chunks.append(text)

    return "".join(chunks), run_spans


def _run_index_for_position(position: int, run_spans: list[tuple[int, int, int]]) -> int | None:
    for run_index, start, end in run_spans:
        if start <= position < end:
            return run_index
    return None


def _find_unbalanced_brackets(full_text: str) -> list[tuple[str, int, int, str]]:
    issues: list[tuple[str, int, int, str]] = []
    open_positions: list[int] = []

    for index, char in enumerate(full_text):
        if char == _OPEN_BRACKET:
            open_positions.append(index)
            continue

        if char == _CLOSE_BRACKET:
            if open_positions:
                open_positions.pop()
            else:
                issues.append(("stray_close", index, index + 1, _CLOSE_BRACKET))

    for start in open_positions:
        issues.append(("unclosed_bracket", start, len(full_text), full_text[start:]))

    return issues
